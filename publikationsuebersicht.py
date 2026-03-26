#!/usr/bin/env python3
"""
Publikationsübersicht — Klinik III für Innere Medizin, Uniklinik Köln
=====================================================================
Automatisches Tool zur Erfassung aller Publikationen der Klinikmitarbeiter.
Scrapt die aktuelle Mitarbeiterliste von der Klinik-Website, durchsucht
PubMed und Google Scholar, bereinigt Duplikate und erstellt ein Word-Dokument.

Nutzung:
    python publikationsuebersicht.py
    python publikationsuebersicht.py --start-date 2024-06-01 --end-date 2025-03-25
    python publikationsuebersicht.py --output mein_bericht.docx
    python publikationsuebersicht.py --no-scholar  (nur PubMed, kein Google Scholar)
"""

import argparse
import re
import sys
import time
import unicodedata
from datetime import datetime, date
from difflib import SequenceMatcher

import requests
from bs4 import BeautifulSoup
from Bio import Entrez, Medline
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# ---------------------------------------------------------------------------
# Konfiguration
# ---------------------------------------------------------------------------
Entrez.email = "publikationen@uk-koeln.de"  # PubMed verlangt eine E-Mail
PUBMED_DELAY = 0.34  # Sekunden zwischen PubMed-Anfragen (max 3/s)
SCHOLAR_DELAY = 5     # Sekunden zwischen Google Scholar-Anfragen

URLS = {
    "direktion": "https://kardiologie.uk-koeln.de/klinik/direktor-team/",
    "oberaerzte": "https://kardiologie.uk-koeln.de/klinik/direktor-team/oberaerzte/",
    "fach_und_assistenz": "https://kardiologie.uk-koeln.de/klinik/direktor-team/fach-und-assistenzaerzte/",
}

AFFILIATION_KEYWORDS = [
    "cologne", "koln", "köln", "koeln",
    "university hospital cologne",
    "universitätsklinikum köln", "uniklinik köln",
    "heart center cologne", "herzzentrum köln",
]


# ---------------------------------------------------------------------------
# 1. Website-Scraper
# ---------------------------------------------------------------------------
def normalize_name(name: str) -> str:
    """Entfernt akademische Titel und normalisiert den Namen."""
    titles = [
        r"Univ\.-Prof\.\s*",
        r"Prof\.\s*",
        r"Priv\.-Doz\.\s*",
        r"PD\s+",
        r"Dr\.\s*med\.\s*",
        r"Dr\.\s*rer\.\s*nat\.\s*",
        r"Dr\.\s*",
        r"PhD\s*",
        r",\s*PhD",
        r",\s*MHBA",
        r"LL\.M\.\s*",
        r",\s*LL\.M\.",
        r"Dipl\.-\w+\s*",
        r"em\.\s*",
    ]
    cleaned = name.strip()
    for t in titles:
        cleaned = re.sub(t, "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    # Trailing-Komma entfernen (z.B. nach Entfernung von ", PhD")
    cleaned = cleaned.rstrip(",").strip()
    return cleaned


def extract_title(name: str) -> str:
    """Extrahiert den akademischen Titel."""
    title_patterns = [
        (r"Univ\.-Prof\.", "Univ.-Prof."),
        (r"Prof\.", "Prof."),
        (r"Priv\.-Doz\.", "Priv.-Doz."),
        (r"Dr\.\s*med\.", "Dr. med."),
        (r"Dr\.", "Dr."),
    ]
    for pattern, label in title_patterns:
        if re.search(pattern, name):
            return label
    return ""


def _scrape_person_cards(soup, category: str) -> list[dict]:
    """Extrahiert Personen aus div.tx-bra-persons-ukk Containern."""
    persons = []
    cards = soup.select("div.tx-bra-persons-ukk")
    for card in cards:
        text = card.get_text(separator="\n").strip()
        lines = [l.strip() for l in text.split("\n") if l.strip()]
        if not lines:
            continue
        full_name = lines[0].strip()
        # Ignoriere Navigationseintraege oder zu kurze Strings
        if len(full_name) < 4 or full_name.startswith("Klinik"):
            continue
        clean = normalize_name(full_name)
        parts = clean.split()
        if len(parts) < 2:
            continue

        # Position/Schwerpunkt extrahieren
        position = ""
        for line in lines[1:5]:
            if any(kw in line for kw in ["Schwerpunkt", "Direktor", "Leiter",
                                          "Stellvertr", "Geschaeftsf", "Geschäftsf",
                                          "Personal", "Chest Pain", "Intermediate",
                                          "Normalstation", "Elektrophysiologie",
                                          "Pneumologie", "Intensivstation"]):
                position = line.strip()[:150]
                break

        persons.append({
            "full_name": full_name,
            "clean_name": clean,
            "title": extract_title(full_name),
            "position": position,
            "category": category,
        })
    return persons


def _scrape_people_teasers(soup, default_category: str) -> list[dict]:
    """Extrahiert Personen aus li.people-teaser__item Eintraegen."""
    persons = []

    # Bestimme Sektionsgrenzen anhand von h2-Headern
    # z.B. "Fachärztinnen und -ärzte" vs "Assistenzärztinnen und -ärzte"
    main = soup.find("main") or soup
    all_elements = []
    for el in main.descendants:
        if hasattr(el, 'name') and el.name in ('h2', 'li'):
            if el.name == 'h2':
                all_elements.append(('header', el.get_text(strip=True)))
            elif 'people-teaser__item' in (el.get('class') or []):
                text = el.get_text(separator="\n").strip()
                first_line = text.split("\n")[0].strip() if text else ""
                all_elements.append(('person', first_line, text))

    current_section = default_category
    for item in all_elements:
        if item[0] == 'header':
            header_text = item[1].lower()
            if "fachärzt" in header_text:
                current_section = "Fachärzte"
            elif "assistenzärzt" in header_text:
                current_section = "Assistenzärzte"
        elif item[0] == 'person':
            full_name = item[1]
            full_text = item[2]
            if len(full_name) < 3:
                continue
            clean = normalize_name(full_name)
            parts = clean.split()
            if len(parts) < 2:
                continue

            # Position extrahieren
            position = ""
            text_lines = [l.strip() for l in full_text.split("\n") if l.strip()]
            for line in text_lines[1:4]:
                if any(kw in line for kw in ["Schwerpunkt", "SFB", "Myeloide",
                                              "Publikationen"]):
                    if "Publikationen" not in line:
                        position = line.strip()[:150]
                    break

            persons.append({
                "full_name": full_name,
                "clean_name": clean,
                "title": extract_title(full_name),
                "position": position,
                "category": current_section,
            })

    return persons


def scrape_staff() -> list[dict]:
    """Scrapt die Mitarbeiterliste von der Klinik-Website."""
    staff = []
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
    }

    page_configs = [
        ("Direktion", URLS["direktion"]),
        ("Oberärzte", URLS["oberaerzte"]),
        ("Fach- und Assistenzärzte", URLS["fach_und_assistenz"]),
    ]

    for category, url in page_configs:
        print(f"  Scrape {category}...")
        try:
            resp = requests.get(url, headers=headers, timeout=30)
            resp.raise_for_status()
            soup = BeautifulSoup(resp.text, "html.parser")

            # Methode 1: Person-Cards (div.tx-bra-persons-ukk)
            # Für Fach-/Assistenzärzte müssen wir die Kategorie aus dem Kontext bestimmen
            cards = _scrape_person_cards(soup, category)
            if category == "Fach- und Assistenzärzte":
                # Bestimme Unterkategorie aus dem HTML-Kontext
                all_text = soup.get_text()
                for card in cards:
                    idx = all_text.find(card["full_name"])
                    if idx > 0:
                        pre = all_text[:idx].lower()
                        if "assistenzärzt" in pre[max(0, len(pre)-1000):]:
                            card["category"] = "Assistenzärzte"
                        elif "fachärzt" in pre[max(0, len(pre)-1000):]:
                            card["category"] = "Fachärzte"
            staff.extend(cards)
            if cards:
                print(f"    {len(cards)} aus Person-Cards")

            # Methode 2: People-Teasers (li.people-teaser__item)
            teasers = _scrape_people_teasers(soup, category)
            staff.extend(teasers)
            if teasers:
                print(f"    {len(teasers)} aus People-Teasers")

        except Exception as e:
            print(f"  WARNUNG: {category} konnte nicht gescrapt werden: {e}")

    # Dedupliziere nach clean_name (case-insensitive)
    seen = set()
    unique_staff = []
    for s in staff:
        key = s["clean_name"].lower()
        if key not in seen:
            seen.add(key)
            unique_staff.append(s)

    # Zusätzliche Mitarbeiter die nicht auf der Website stehen
    EXTRA_STAFF = [
        {"full_name": "Per Arkenberg", "clean_name": "Per Arkenberg",
         "title": "", "position": "", "category": "Mitarbeiter"},
        {"full_name": "Harshal Nemade", "clean_name": "Harshal Nemade",
         "title": "", "position": "", "category": "Mitarbeiter"},
        {"full_name": "Kristel Martinez Lagunas", "clean_name": "Kristel Martinez Lagunas",
         "title": "", "position": "", "category": "Mitarbeiter"},
        {"full_name": "Yein Park", "clean_name": "Yein Park",
         "title": "", "position": "", "category": "Mitarbeiter"},
        {"full_name": "Elvina Santhamma Philip", "clean_name": "Elvina Santhamma Philip",
         "title": "", "position": "", "category": "Mitarbeiter"},
        {"full_name": "Suchitra Narayan", "clean_name": "Suchitra Narayan",
         "title": "", "position": "", "category": "Mitarbeiter"},
        {"full_name": "Chantal Wientjens", "clean_name": "Chantal Wientjens",
         "title": "", "position": "", "category": "Mitarbeiter"},
        {"full_name": "Holger Winkels", "clean_name": "Holger Winkels",
         "title": "", "position": "", "category": "Mitarbeiter"},
        {"full_name": "Martin Mollenhauer", "clean_name": "Martin Mollenhauer",
         "title": "", "position": "", "category": "Mitarbeiter"},
        {"full_name": "Valerie Lohner", "clean_name": "Valerie Lohner",
         "title": "", "position": "", "category": "Mitarbeiter"},
        {"full_name": "Thomas Riffelmacher", "clean_name": "Thomas Riffelmacher",
         "title": "", "position": "", "category": "Mitarbeiter"},
        {"full_name": "Arian Sultan", "clean_name": "Arian Sultan",
         "title": "", "position": "", "category": "Ehemaliger Oberarzt"},
        {"full_name": "Simon Braumann", "clean_name": "Simon Braumann",
         "title": "", "position": "", "category": "Ehemaliger Oberarzt"},
        {"full_name": "Christoph Adler", "clean_name": "Christoph Adler",
         "title": "", "position": "", "category": "Ehemaliger Oberarzt"},
    ]

    for extra in EXTRA_STAFF:
        key = extra["clean_name"].lower()
        if key not in seen:
            seen.add(key)
            unique_staff.append(extra)

    return unique_staff


# ---------------------------------------------------------------------------
# 2. PubMed-Suche
# ---------------------------------------------------------------------------
def _umlaut_variants(name: str) -> list[str]:
    """Gibt Umlaut-Varianten eines Namens zurück.

    Erzeugt drei Varianten:
      1. Original
      2. Umlaute -> Digraphen (ä->ae, ö->oe, ü->ue, ß->ss)
      3. Umlaute -> einfacher Vokal (ä->a, ö->o, ü->u)  ← PubMed nutzt oft diese!
      4. Digraphen -> Umlaute (ae->ä, oe->ö, ue->ü)
    """
    umlaut_to_digraph = {"ä": "ae", "ö": "oe", "ü": "ue", "ß": "ss"}
    umlaut_to_simple = {"ä": "a", "ö": "o", "ü": "u", "ß": "ss"}
    digraph_to_umlaut = {"ae": "ä", "oe": "ö", "ue": "ü"}
    variants = {name}

    # Umlaute -> Digraphen (ü -> ue)
    alt = name
    for k, v in umlaut_to_digraph.items():
        alt = alt.replace(k, v)
    if alt != name:
        variants.add(alt)

    # Umlaute -> einfacher Vokal (ü -> u)  — PubMed-Standard!
    alt2 = name
    for k, v in umlaut_to_simple.items():
        alt2 = alt2.replace(k, v)
    if alt2 != name:
        variants.add(alt2)

    # Digraphen -> Umlaute (ue -> ü)
    alt3 = name
    for k, v in digraph_to_umlaut.items():
        alt3 = alt3.replace(k, v)
    if alt3 != name:
        variants.add(alt3)

    return list(variants)


# Namenspräfixe die zum Nachnamen gehören
NAME_PREFIXES = {"ten", "van", "von", "de", "di", "del", "den", "der", "la", "le"}


def _split_name(clean_name: str) -> tuple[str, list[str], str]:
    """Teilt einen bereinigten Namen in (Vorname, Mittelnamen, Nachname).
    Beachtet Namenspräfixe wie 'ten', 'van den', 'von', 'di'."""
    parts = clean_name.split()
    if len(parts) < 2:
        return (clean_name, [], "")

    first = parts[0]

    # Finde wo der Nachname beginnt (inkl. Präfixe)
    # Gehe von hinten und prüfe ob es ein Präfix gibt
    last_start = len(parts) - 1
    while last_start > 1 and parts[last_start - 1].lower() in NAME_PREFIXES:
        last_start -= 1
    # Mindestens 1 Wort als Vorname
    if last_start < 1:
        last_start = 1

    middle = parts[1:last_start]
    last_name = " ".join(parts[last_start:])

    return (first, middle, last_name)


def make_pubmed_queries(person: dict) -> list[str]:
    """Erzeugt verschiedene PubMed-Suchvarianten für eine Person."""
    clean = person["clean_name"]
    first, middle, last_name = _split_name(clean)
    if not last_name:
        return []

    queries = []

    # Alle Nachname-Varianten (mit/ohne Umlaute)
    last_variants = _umlaut_variants(last_name)
    first_variants = _umlaut_variants(first)

    for lv in last_variants:
        for fv in first_variants:
            # "Nachname Vorname"[Author]
            queries.append(f'"{lv} {fv}"[Author]')
            # "Nachname V"[Author] (Initial)
            queries.append(f'"{lv} {fv[0]}"[Author]')

        # Wenn es Mittelnamen gibt: "Nachname VI" (beide Initialen)
        if middle:
            initials = first[0] + "".join(m[0] for m in middle)
            queries.append(f'"{lv} {initials}"[Author]')
            # Auch nur mit Mittel-Initial
            for mv in middle:
                queries.append(f'"{lv} {first[0]}{mv[0]}"[Author]')

    return list(dict.fromkeys(queries))  # Dedupliziere, Reihenfolge beibehalten


def search_pubmed(person: dict, start_date: str, end_date: str) -> list[dict]:
    """Sucht PubMed-Artikel für eine Person im gegebenen Zeitraum."""
    queries = make_pubmed_queries(person)
    all_pmids = set()
    articles = []

    # Zwei Durchläufe: erst MIT Affiliation-Filter, dann OHNE (für Fälle wo
    # Affiliation nicht indexiert ist, aber nur für die ersten Query-Varianten)
    date_filter = f'("{start_date}"[Date - Publication] : "{end_date}"[Date - Publication])'
    affil_filter = '(Cologne[Affiliation] OR Koln[Affiliation] OR Köln[Affiliation])'

    search_rounds = []
    # Runde 1: Mit Affiliation-Filter (alle Query-Varianten)
    for query in queries:
        search_rounds.append(f'{query} AND {affil_filter} AND {date_filter}')

    # Runde 2: Nur Nachname + Affiliation (fängt Autoren mit unbekannten
    # Mittelnamen-Initialen auf, z.B. "Nettersheim FS" statt "Nettersheim F")
    first, middle, last_name = _split_name(person["clean_name"])
    if last_name:
        last_variants = _umlaut_variants(last_name)
        for lv in last_variants:
            search_rounds.append(
                f'{lv}[Author] AND {affil_filter} AND {date_filter}'
            )

    # Runde 3: Voller Name (ohne [Author]-Tag) + Affiliation
    # PubMed durchsucht dann ALLE Felder — findet auch Autoren mit
    # ungewöhnlichen Initialen oder abweichenden Namensformaten.
    clean = person["clean_name"]
    search_rounds.append(
        f'{clean} AND {affil_filter} AND {date_filter}'
    )

    # Runde 4: Spezifischste Query OHNE Affiliation-Filter
    # Fängt Artikel auf, bei denen PubMed keine Affiliations indexiert hat
    # (z.B. JACC Letters, Epub ahead of print).
    # Die strikte FAU-Prüfung (SequenceMatcher > 0.7) + common_short_names
    # verhindert False Positives.
    if queries:
        search_rounds.append(f'{queries[0]} AND {date_filter}')

    for full_query in search_rounds:
        try:
            time.sleep(PUBMED_DELAY)
            handle = Entrez.esearch(db="pubmed", term=full_query, retmax=200)
            results = Entrez.read(handle)
            handle.close()
            pmids = results.get("IdList", [])
            new_pmids = [p for p in pmids if p not in all_pmids]
            all_pmids.update(pmids)

            if not new_pmids:
                continue

            # Fetch in Batches von max 50
            for batch_start in range(0, len(new_pmids), 50):
                batch = new_pmids[batch_start:batch_start+50]
                time.sleep(PUBMED_DELAY)
                handle = Entrez.efetch(
                    db="pubmed", id=",".join(batch),
                    rettype="medline", retmode="text"
                )
                records = list(Medline.parse(handle))
                handle.close()

                for rec in records:
                    pmid = rec.get("PMID", "")
                    if any(a["pmid"] == pmid for a in articles):
                        continue

                    # Affiliation-Prüfung
                    affiliations = rec.get("AD", "")
                    if isinstance(affiliations, list):
                        affiliations = " ".join(affiliations)

                    has_cologne = any(
                        kw in affiliations.lower()
                        for kw in AFFILIATION_KEYWORDS
                    )

                    # Autorenposition bestimmen
                    authors = rec.get("AU", [])
                    full_authors = rec.get("FAU", [])
                    author_position = determine_author_position(
                        person, authors, full_authors
                    )

                    # Prüfe auch Investigator-Liste (FIR/IR)
                    is_investigator = False
                    if author_position == "nicht gefunden":
                        full_investigators = rec.get("FIR", [])
                        investigators = rec.get("IR", [])
                        if full_investigators or investigators:
                            inv_pos = determine_author_position(
                                person, investigators, full_investigators
                            )
                            if inv_pos != "nicht gefunden":
                                is_investigator = True
                                author_position = "Investigator"

                    # Strikte Filterung:
                    # - Muss als Autor oder Investigator gefunden werden UND
                    # - Muss Köln-Affiliation haben (oder der Artikel kam aus
                    #   der Affiliation-gefilterten Suche)
                    if author_position == "nicht gefunden":
                        continue

                    # Wenn keine Köln-Affiliation gefunden: nur akzeptieren
                    # wenn der Name sehr spezifisch ist (>= 3 Namensbestandteile
                    # oder ungewöhnlicher Nachname)
                    if not has_cologne:
                        _, _, p_last = _split_name(person["clean_name"])
                        last_name = p_last.split()[-1] if p_last else ""
                        # Häufige kurze Nachnamen brauchen zwingend Affiliation
                        common_short_names = {
                            "lee", "adam", "frank", "ernst", "marx", "weber",
                            "schmidt", "meyer", "mueller", "muller", "fischer",
                            "hartmann", "wagner", "becker", "schulz", "hoffmann",
                            "schaefer", "schäfer", "koch", "richter", "klein",
                            "wolf", "schröder", "neumann", "schwarz", "zimmermann",
                            "braun", "krüger", "hofmann", "lange", "schmitt",
                            "werner", "krause", "meier", "lehmann", "huber",
                            "mayer", "herrmann", "könig", "walter", "kaiser",
                            "fuchs", "peters", "lang", "scholz", "jung",
                            "hahn", "keller", "vogel", "schubert", "roth",
                            "stein", "wilson", "pfeiffer", "jahn", "rath",
                            "wrobel", "finke", "adler", "dittrich", "park",
                            "wang", "chen", "liu", "zhang", "li", "kim",
                            "hohmann", "rosenkranz", "hoyer", "potthoff",
                            "grossmann", "großmann", "koerber", "körber",
                            "di benedetto", "lohner", "narayan",
                        }
                        if last_name.lower() in common_short_names:
                            continue
                        # Für alle anderen: ohne Affiliation trotzdem riskant
                        # bei sehr kurzen Nachnamen (< 5 Zeichen).
                        # Das strikte FAU-Matching (SequenceMatcher > 0.7)
                        # verhindert False Positives auch bei längeren Namen.
                        name_parts = person["clean_name"].split()
                        if len(name_parts) < 3 and len(last_name) < 5:
                            continue

                    # Datum parsen
                    pub_date = rec.get("DP", "")
                    parsed_date = parse_pubmed_date(pub_date)

                    doi = ""
                    aid = rec.get("AID", [])
                    if isinstance(aid, list):
                        for a in aid:
                            if a.endswith("[doi]"):
                                doi = a.replace(" [doi]", "").strip()
                                break
                    elif isinstance(aid, str) and "[doi]" in aid:
                        doi = aid.replace(" [doi]", "").strip()

                    articles.append({
                        "pmid": pmid,
                        "title": rec.get("TI", ""),
                        "authors": authors,
                        "full_authors": full_authors,
                        "journal": rec.get("JT", rec.get("TA", "")),
                        "date": parsed_date,
                        "date_raw": pub_date,
                        "doi": doi,
                        "author_position": author_position,
                        "source": "PubMed",
                        "has_cologne_affiliation": has_cologne,
                        "assigned_to": person["clean_name"],
                        "assigned_category": person["category"],
                        "assigned_full_name": person["full_name"],
                    })

        except Exception as e:
            print(f"    PubMed-Fehler bei Query '{full_query[:80]}...': {e}")

    return articles


def determine_author_position(person: dict, authors: list[str],
                               full_authors: list[str] | None = None) -> str:
    """Bestimmt die Autorenposition einer Person in der Autorenliste.

    Nutzt bevorzugt full_authors (FAU-Feld, z.B. 'Schaefer, Matthieu')
    für präzises Matching mit vollem Vornamen. Fällt auf authors (AU-Feld,
    z.B. 'Schaefer M') zurück, wenn FAU nicht verfügbar.
    """
    first, middle, last_name = _split_name(person["clean_name"])
    if not last_name:
        return "nicht gefunden"

    first_lower = first.lower()
    first_initial = first[0].lower()

    # Alle Nachname-Varianten (mit/ohne Umlaute)
    last_variants = [v.lower() for v in _umlaut_variants(last_name)]
    # Alle Vorname-Varianten
    first_variants = [v.lower() for v in _umlaut_variants(first)]

    num_authors = len(full_authors) if full_authors else len(authors)

    def _position_label(idx: int) -> str:
        if idx == 0:
            return "Erstautor"
        elif idx == num_authors - 1:
            return "Letztautor"
        elif idx == num_authors - 2:
            return "Vorletzter Autor"
        else:
            return f"Ko-Autor (Position {idx+1}/{num_authors})"

    # --- Versuch 1: Voller Name über FAU-Feld (z.B. "Schaefer, Matthieu") ---
    if full_authors:
        for i, fau in enumerate(full_authors):
            fau_lower = fau.lower()
            # FAU-Format: "Nachname, Vorname Mittelname"
            last_match = any(v in fau_lower for v in last_variants)
            if not last_match:
                continue

            # Prüfe vollen Vornamen mit hoher Strenge
            fau_parts = fau.split(",", 1)
            if len(fau_parts) >= 2:
                fau_first_part = fau_parts[1].strip().lower()
                fau_first_names = fau_first_part.split()
                fau_first = fau_first_names[0] if fau_first_names else ""

                first_match = False
                for fv in first_variants:
                    if len(fv) >= 4 and len(fau_first) >= 4:
                        # Beide Namen lang genug: SequenceMatcher > 0.7
                        ratio = SequenceMatcher(None, fv, fau_first).ratio()
                        if ratio > 0.7:
                            first_match = True
                            break
                    elif len(fv) >= 2 and len(fau_first) >= 2:
                        # Kurze Namen: exakter Match erforderlich
                        if fv == fau_first:
                            first_match = True
                            break
                    elif fv and fau_first and fv[0] == fau_first[0]:
                        first_match = True
                        break

                if first_match:
                    return _position_label(i)

    # --- Versuch 2: Kurzformat AU-Feld (z.B. "Schaefer M") als Fallback ---
    # NUR wenn FAU NICHT verfügbar war! Wenn FAU vorhanden ist und keinen
    # Match ergab, ist das eine definitive Ablehnung (z.B. Christian ≠ Christopher).
    if not full_authors:
        for i, author in enumerate(authors):
            author_lower = author.lower()
            name_match = any(v in author_lower for v in last_variants)
            if not name_match:
                continue

            author_parts = author.split()
            if len(author_parts) >= 2:
                author_initials = author_parts[-1].lower() if len(author_parts[-1]) <= 3 else ""
                if first_initial in author_initials or first[:3].lower() in author_lower:
                    return _position_label(i)

            # Wenn nur Nachname passt und Nachname lang/einzigartig genug
            if len(last_name) >= 8:
                return _position_label(i)

    return "nicht gefunden"


def parse_pubmed_date(date_str: str) -> date:
    """Parst ein PubMed-Datum in ein date-Objekt."""
    if not date_str:
        return date(1900, 1, 1)
    # Typische Formate: "2025 Mar 15", "2025 Mar", "2025", "2025 Spring"
    date_str = date_str.strip()
    seasons = {"spring": "03", "summer": "06", "fall": "09", "autumn": "09", "winter": "12"}

    for season, month in seasons.items():
        if season in date_str.lower():
            year = re.search(r"(\d{4})", date_str)
            if year:
                return date(int(year.group(1)), int(month), 1)

    formats = ["%Y %b %d", "%Y %b", "%Y"]
    for fmt in formats:
        try:
            return datetime.strptime(date_str[:len("2025 Mar 15")], fmt).date()
        except ValueError:
            continue

    year_match = re.search(r"(\d{4})", date_str)
    if year_match:
        return date(int(year_match.group(1)), 1, 1)

    return date(1900, 1, 1)


# ---------------------------------------------------------------------------
# 3. Google Scholar-Suche (optional)
# ---------------------------------------------------------------------------
def search_scholar(person: dict, start_year: int, end_year: int) -> list[dict]:
    """Sucht Google Scholar nach Publikationen einer Person."""
    try:
        from scholarly import scholarly as scholar_api
    except ImportError:
        print("    scholarly nicht installiert — überspringe Google Scholar")
        return []

    clean = person["clean_name"]
    parts = clean.split()
    if len(parts) < 2:
        return []

    articles = []
    query = f'author:"{parts[-1]}" "{clean}" cardiology Cologne'

    try:
        time.sleep(SCHOLAR_DELAY)
        search_results = scholar_api.search_pubs(query, year_low=start_year, year_high=end_year)

        count = 0
        for result in search_results:
            if count >= 20:  # Begrenze pro Person
                break
            count += 1

            bib = result.get("bib", {})
            title = bib.get("title", "")
            authors_str = bib.get("author", "")
            if isinstance(authors_str, list):
                authors = authors_str
            else:
                authors = [a.strip() for a in authors_str.split(" and ")]

            pub_year = bib.get("pub_year", "")
            journal = bib.get("venue", bib.get("journal", ""))

            # Prüfe ob diese Person wirklich Autor ist
            person_in_authors = False
            last_name = parts[-1].lower()
            for a in authors:
                if last_name in a.lower():
                    person_in_authors = True
                    break

            if not person_in_authors:
                continue

            # DOI extrahieren falls vorhanden
            doi = ""
            pub_url = result.get("pub_url", "")
            if "doi.org" in pub_url:
                doi = pub_url.split("doi.org/")[-1]

            try:
                parsed_date = date(int(pub_year), 1, 1) if pub_year else date(1900, 1, 1)
            except (ValueError, TypeError):
                parsed_date = date(1900, 1, 1)

            # Autorenposition
            position = "Ko-Autor"
            if authors:
                first_author_last = authors[0].split()[-1].lower() if authors[0] else ""
                last_author_last = authors[-1].split()[-1].lower() if authors[-1] else ""
                if last_name == first_author_last:
                    position = "Erstautor"
                elif last_name == last_author_last:
                    position = "Letztautor"

            articles.append({
                "pmid": "",
                "title": title,
                "authors": authors,
                "full_authors": authors,  # Scholar hat schon volle Namen
                "journal": journal,
                "date": parsed_date,
                "date_raw": pub_year,
                "doi": doi,
                "author_position": position,
                "source": "Google Scholar",
                "has_cologne_affiliation": True,  # Kann nicht geprüft werden
                "assigned_to": person["clean_name"],
                "assigned_category": person["category"],
                "assigned_full_name": person["full_name"],
            })

            time.sleep(1)  # Kurze Pause zwischen Ergebnissen

    except Exception as e:
        print(f"    Google Scholar-Fehler für {clean}: {e}")

    return articles


# ---------------------------------------------------------------------------
# 4. Duplikat-Erkennung
# ---------------------------------------------------------------------------
def normalize_title(title: str) -> str:
    """Normalisiert einen Titel für den Vergleich."""
    # Kleinbuchstaben, Sonderzeichen entfernen
    title = title.lower().strip()
    title = re.sub(r"[^\w\s]", "", title)
    title = re.sub(r"\s+", " ", title)
    return title


def are_duplicates(art1: dict, art2: dict) -> bool:
    """Prüft ob zwei Artikel Duplikate sind."""
    # DOI-Match
    if art1.get("doi") and art2.get("doi"):
        doi1 = art1["doi"].lower().strip().rstrip(".")
        doi2 = art2["doi"].lower().strip().rstrip(".")
        if doi1 == doi2:
            return True

    # PMID-Match
    if art1.get("pmid") and art2.get("pmid"):
        if art1["pmid"] == art2["pmid"]:
            return True

    # Titel-Ähnlichkeit
    t1 = normalize_title(art1.get("title", ""))
    t2 = normalize_title(art2.get("title", ""))
    if t1 and t2:
        ratio = SequenceMatcher(None, t1, t2).ratio()
        if ratio > 0.85:
            return True

    return False


POSITION_RANK = {
    "Erstautor": 1,
    "Letztautor": 2,
    "Vorletzter Autor": 3,
}


def get_position_rank(position: str) -> int:
    """Gibt den Rang einer Autorenposition zurück (kleiner = besser)."""
    if position in POSITION_RANK:
        return POSITION_RANK[position]
    if position.startswith("Ko-Autor"):
        return 10
    return 99


def deduplicate_articles(all_articles: list[dict]) -> list[dict]:
    """Bereinigt Duplikate und ordnet jedem Artikel den signifikantesten Autor zu."""
    if not all_articles:
        return []

    # Gruppiere Duplikate
    groups = []
    used = [False] * len(all_articles)

    for i, art in enumerate(all_articles):
        if used[i]:
            continue
        group = [art]
        used[i] = True
        for j in range(i + 1, len(all_articles)):
            if used[j]:
                continue
            if are_duplicates(art, all_articles[j]):
                group.append(all_articles[j])
                used[j] = True
        groups.append(group)

    # Pro Gruppe: wähle den besten Eintrag und vermerke Co-Autoren
    result = []
    for group in groups:
        # Corrections/Errata NICHT als Duplikate zusammenfassen —
        # sowohl Original als auch Correction separat auflisten
        correction_keywords = ["correction:", "erratum", "corrigendum", "publisher correction"]
        corrections = [a for a in group
                       if any(kw in a.get("title", "").lower() for kw in correction_keywords)]
        originals = [a for a in group
                     if not any(kw in a.get("title", "").lower() for kw in correction_keywords)]
        if corrections and originals:
            # Corrections als eigene Einträge behalten
            for corr in corrections:
                corr_entry = dict(corr)
                corr_entry["other_clinic_authors"] = []
                corr_entry["all_clinic_authors"] = [corr["assigned_to"]]
                result.append(corr_entry)
            group = originals  # Weiter mit nur den Originalen

        # Sortiere nach Autorenposition (signifikanteste zuerst)
        group.sort(key=lambda x: get_position_rank(x["author_position"]))
        best = dict(group[0])  # Kopie

        # Sammle ALLE Klinik-Mitarbeiter die auf diesem Artikel stehen
        all_clinic_authors = {best["assigned_to"].lower(): best["assigned_to"]}
        for art in group[1:]:
            key = art["assigned_to"].lower()
            if key not in all_clinic_authors:
                all_clinic_authors[key] = art["assigned_to"]

        best["all_clinic_authors"] = list(all_clinic_authors.values())

        # Legacy-Feld beibehalten für Kompatibilität
        other_clinic_authors = []
        for art in group[1:]:
            if art["assigned_to"] != best["assigned_to"]:
                other_clinic_authors.append({
                    "name": art["assigned_full_name"],
                    "position": art["author_position"],
                    "category": art["assigned_category"],
                })
        best["other_clinic_authors"] = other_clinic_authors

        # Bevorzuge PubMed-Daten (vollständiger)
        pubmed_entry = None
        for art in group:
            if art["source"] == "PubMed":
                pubmed_entry = art
                break
        if pubmed_entry and best["source"] != "PubMed":
            # Übernehme PubMed-Daten, behalte aber die Zuordnung
            assigned = best["assigned_to"]
            assigned_full = best["assigned_full_name"]
            assigned_cat = best["assigned_category"]
            assigned_pos = best["author_position"]
            others = best["other_clinic_authors"]
            best.update(pubmed_entry)
            best["assigned_to"] = assigned
            best["assigned_full_name"] = assigned_full
            best["assigned_category"] = assigned_cat
            best["author_position"] = assigned_pos
            best["other_clinic_authors"] = others

        result.append(best)

    # Sortiere nach Datum (neueste zuerst)
    result.sort(key=lambda x: x["date"], reverse=True)

    return result


# ---------------------------------------------------------------------------
# 5. Word-Dokument-Generator
# ---------------------------------------------------------------------------
def create_word_document(
    articles: list[dict],
    staff: list[dict],
    start_date: str,
    end_date: str,
    output_path: str,
):
    """Erstellt das Word-Dokument mit der Publikationsübersicht."""
    doc = Document()

    # Seitenränder
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Titel
    title = doc.add_heading(
        "Publikationsübersicht", level=0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading(
        "Klinik III für Innere Medizin — Kardiologie", level=1
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadaten
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Zähle Erst-/Letztautorschaften für die Metadaten
    n_first = sum(1 for a in articles if a.get("author_position") == "Erstautor")
    n_last = sum(1 for a in articles if a.get("author_position") == "Letztautor")
    meta_run = meta.add_run(
        f"Zeitraum: {start_date} bis {end_date}\n"
        f"Erstellt am: {datetime.now().strftime('%d.%m.%Y %H:%M')}\n"
        f"Anzahl Mitarbeiter durchsucht: {len(staff)}\n"
        f"Anzahl Publikationen: {len(articles)}\n"
        f"Davon Erstautorschaften: {n_first} | Letztautorschaften: {n_last}"
    )
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()  # Leerzeile

    if not articles:
        doc.add_paragraph(
            "Keine Publikationen im angegebenen Zeitraum gefunden."
        )
        doc.save(output_path)
        return

    # Zusammenfassung nach Mitarbeiter
    doc.add_heading("Zusammenfassung nach Mitarbeiter", level=2)

    def _get_real_position(fau_name: str, authors_list) -> str:
        """Bestimmt die echte Position eines Autors in der Autorenliste."""
        if not authors_list or not isinstance(authors_list, list):
            return "Ko-Autor"
        # Finde den Autor in der Liste
        fau_lower = fau_name.lower().strip()
        for idx, author in enumerate(authors_list):
            author_lower = author.lower().strip()
            # Prüfe ob der FAU-Name im Autornamen enthalten ist
            if fau_lower in author_lower or author_lower in fau_lower:
                if idx == 0:
                    return "Erstautor"
                elif idx == len(authors_list) - 1:
                    return "Letztautor"
                else:
                    return "Ko-Autor"
            # Prüfe Nachname-Match
            fau_parts = fau_name.split(",", 1) if "," in fau_name else [fau_name]
            fau_last = fau_parts[0].strip().lower()
            if fau_last and len(fau_last) > 2:
                author_parts = author.split(",", 1) if "," in author else author.split()
                author_last = author_parts[0].strip().lower() if "," in author else (author_parts[-1].strip().lower() if author_parts else "")
                if fau_last == author_last:
                    if idx == 0:
                        return "Erstautor"
                    elif idx == len(authors_list) - 1:
                        return "Letztautor"
                    else:
                        return "Ko-Autor"
        return "Ko-Autor"

    # Zähle Publikationen pro Mitarbeiter — echte Position in der Autorenliste
    pub_counts = {}
    for art in articles:
        authors_list = art.get("full_authors", art.get("authors", []))
        if not isinstance(authors_list, list):
            authors_list = []

        counted_names = set()

        # Zugeordneter Autor: echte Position aus der Autorenliste
        name_key = art["assigned_full_name"]
        pos = art.get("author_position", "Ko-Autor")
        counted_names.add(art["assigned_to"].lower())
        if name_key not in pub_counts:
            pub_counts[name_key] = {"Gesamt": 0, "Erstautor": 0, "Letztautor": 0, "Ko-Autor": 0}
        pub_counts[name_key]["Gesamt"] += 1
        if pos == "Erstautor":
            pub_counts[name_key]["Erstautor"] += 1
        elif pos == "Letztautor":
            pub_counts[name_key]["Letztautor"] += 1
        else:
            pub_counts[name_key]["Ko-Autor"] += 1

        # Andere Klinik-Autoren: echte Position individuell bestimmen
        for o in art.get("other_clinic_authors", []):
            oname = normalize_name(o["name"]).lower()
            if oname not in counted_names:
                counted_names.add(oname)
                # Echte Position aus der Autorenliste bestimmen
                real_pos = _get_real_position(o["name"], authors_list)
                if o["name"] not in pub_counts:
                    pub_counts[o["name"]] = {"Gesamt": 0, "Erstautor": 0, "Letztautor": 0, "Ko-Autor": 0}
                pub_counts[o["name"]]["Gesamt"] += 1
                if real_pos == "Erstautor":
                    pub_counts[o["name"]]["Erstautor"] += 1
                elif real_pos == "Letztautor":
                    pub_counts[o["name"]]["Letztautor"] += 1
                else:
                    pub_counts[o["name"]]["Ko-Autor"] += 1

    # Nur Mitarbeiter mit Publikationen anzeigen
    pub_counts = {k: v for k, v in pub_counts.items() if v["Gesamt"] > 0}

    # Tabelle
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["Mitarbeiter", "Erst", "Letzt", "Ko"]):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)

    for name, counts in sorted(pub_counts.items(), key=lambda x: x[1]["Gesamt"], reverse=True):
        row = table.add_row()
        row.cells[0].text = name
        row.cells[1].text = str(counts["Erstautor"]) if counts["Erstautor"] else ""
        row.cells[2].text = str(counts["Letztautor"]) if counts["Letztautor"] else ""
        row.cells[3].text = str(counts["Ko-Autor"]) if counts["Ko-Autor"] else ""
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_page_break()

    # --- Hilfsfunktion: Prüfe ob ein FAU-Name ein Klinik-Mitarbeiter ist ---
    staff_names_clean = {s["clean_name"].lower() for s in staff}
    # Baue Lookup: Nachname -> set von vollen clean_names
    staff_by_last = {}
    for s in staff:
        parts = s["clean_name"].split()
        if len(parts) >= 2:
            # Berücksichtige Namenspräfixe
            _, _, s_last = _split_name(s["clean_name"])
            for variant in _umlaut_variants(s_last.lower()):
                if variant not in staff_by_last:
                    staff_by_last[variant] = []
                staff_by_last[variant].append(s["clean_name"])

    def _is_clinic_author(fau_name: str) -> bool:
        """Prüft ob ein FAU-Name (z.B. 'von Stein, Philipp') ein Klinikmitarbeiter ist."""
        # Versuche FAU-Format: "Nachname, Vorname"
        if "," in fau_name:
            parts = fau_name.split(",", 1)
            fau_last = parts[0].strip().lower()
            fau_first = parts[1].strip().split()[0].lower() if parts[1].strip() else ""
        else:
            # Kein Komma → versuche "Vorname Nachname"
            parts = fau_name.strip().split()
            fau_first = parts[0].lower() if parts else ""
            fau_last = parts[-1].lower() if len(parts) >= 2 else ""

        # Suche in staff_by_last — prüfe alle Varianten des Nachnamens
        # Auch Teilnamen prüfen (z.B. "von Stein" im FAU vs "stein" im Index)
        fau_last_variants = [v for v in _umlaut_variants(fau_last)]
        # Für zusammengesetzte Nachnamen: auch den letzten Teil prüfen
        fau_last_parts = fau_last.split()
        if len(fau_last_parts) > 1:
            fau_last_variants.extend(_umlaut_variants(fau_last_parts[-1]))

        for variant in fau_last_variants:
            if variant in staff_by_last:
                for clean_name in staff_by_last[variant]:
                    cn_first = clean_name.split()[0].lower()
                    cn_first_variants = [v.lower() for v in _umlaut_variants(cn_first)]
                    for cfv in cn_first_variants:
                        if len(cfv) >= 4 and len(fau_first) >= 4:
                            # Beide Namen lang genug: SequenceMatcher
                            ratio = SequenceMatcher(None, cfv, fau_first).ratio()
                            if ratio > 0.7:
                                return True
                        elif len(fau_first) <= 2 and fau_first:
                            # FAU hat nur Initial(en): prüfe ob Initial passt
                            if cfv and cfv[0] == fau_first[0]:
                                return True
                        elif cfv == fau_first:
                            return True
        return False

    # Publikationsliste nach Datum
    doc.add_heading("Publikationen", level=2)

    for i, art in enumerate(articles, 1):
        # Nummer und Titel
        p_title = doc.add_paragraph()
        num_run = p_title.add_run(f"{i}. ")
        num_run.bold = True
        num_run.font.size = Pt(10)
        title_run = p_title.add_run(art["title"])
        title_run.bold = True
        title_run.font.size = Pt(10)

        # Autorenliste mit Klinik-Autoren FETT
        details = doc.add_paragraph()
        details.paragraph_format.space_before = Pt(2)
        details.paragraph_format.space_after = Pt(2)

        display_authors = art.get("full_authors", art["authors"])
        if not display_authors:
            display_authors = art["authors"]

        if isinstance(display_authors, list):
            for j, author in enumerate(display_authors):
                is_clinic = _is_clinic_author(author)
                r = details.add_run(author)
                r.font.size = Pt(9)
                if is_clinic:
                    r.bold = True
                if j < len(display_authors) - 1:
                    r = details.add_run("; ")
                    r.font.size = Pt(9)
        else:
            r = details.add_run(str(display_authors))
            r.font.size = Pt(9)

        # Journal, Datum, DOI, PMID
        meta_line = doc.add_paragraph()
        meta_line.paragraph_format.space_before = Pt(1)
        meta_line.paragraph_format.space_after = Pt(4)

        r = meta_line.add_run(art["journal"])
        r.font.size = Pt(9)
        r.italic = True

        r = meta_line.add_run(f"  ({art['date_raw']})")
        r.font.size = Pt(9)

        if art.get("doi"):
            r = meta_line.add_run(f"  |  DOI: {art['doi']}")
            r.font.size = Pt(8)
            r.font.color.rgb = RGBColor(80, 80, 80)

        if art.get("pmid"):
            r = meta_line.add_run(f"  |  PMID: {art['pmid']}")
            r.font.size = Pt(8)
            r.font.color.rgb = RGBColor(80, 80, 80)

        # Investigator-Bemerkung
        if art.get("author_position") == "Investigator":
            inv_note = doc.add_paragraph()
            inv_note.paragraph_format.space_before = Pt(1)
            inv_note.paragraph_format.space_after = Pt(2)
            r = inv_note.add_run(
                f"Hinweis: {art['assigned_full_name']} ist bei diesem Artikel "
                f"als Investigator gelistet, nicht als Autor."
            )
            r.font.size = Pt(8)
            r.italic = True
            r.font.color.rgb = RGBColor(120, 120, 120)

        # Trennlinie
        if i < len(articles):
            sep = doc.add_paragraph()
            sep.paragraph_format.space_before = Pt(0)
            sep.paragraph_format.space_after = Pt(2)
            r = sep.add_run("─" * 90)
            r.font.size = Pt(5)
            r.font.color.rgb = RGBColor(200, 200, 200)

    doc.save(output_path)


# ---------------------------------------------------------------------------
# 6. Hauptprogramm
# ---------------------------------------------------------------------------
def main():
    parser = argparse.ArgumentParser(
        description="Publikationsübersicht — Klinik III für Innere Medizin, Uniklinik Köln"
    )
    parser.add_argument(
        "--start-date",
        default="2025-01-01",
        help="Startdatum (YYYY-MM-DD), Standard: 2025-01-01",
    )
    parser.add_argument(
        "--end-date",
        default=datetime.now().strftime("%Y-%m-%d"),
        help="Enddatum (YYYY-MM-DD), Standard: heute",
    )
    parser.add_argument(
        "--output",
        default=None,
        help="Ausgabedatei (.docx), Standard: Publikationsuebersicht_YYYY-MM-DD.docx",
    )
    parser.add_argument(
        "--no-scholar",
        action="store_true",
        help="Google Scholar-Suche überspringen (nur PubMed)",
    )
    parser.add_argument(
        "--email",
        default="publikationen@uk-koeln.de",
        help="E-Mail für PubMed API (Standard: publikationen@uk-koeln.de)",
    )

    args = parser.parse_args()
    Entrez.email = args.email

    if args.output is None:
        args.output = f"Publikationsuebersicht_{datetime.now().strftime('%Y-%m-%d')}.docx"

    # PubMed-Datumsformat
    start_date_pm = args.start_date.replace("-", "/")
    end_date_pm = args.end_date.replace("-", "/")
    start_year = int(args.start_date[:4])
    end_year = int(args.end_date[:4])

    print("=" * 70)
    print("Publikationsübersicht — Klinik III für Innere Medizin")
    print(f"Zeitraum: {args.start_date} bis {args.end_date}")
    print("=" * 70)

    # Schritt 1: Mitarbeiter scrapen
    print("\n[1/4] Scrape Mitarbeiterliste von Website...")
    staff = scrape_staff()
    print(f"  → {len(staff)} Mitarbeiter gefunden")

    if not staff:
        print("FEHLER: Keine Mitarbeiter gefunden. Bitte Internetverbindung prüfen.")
        sys.exit(1)

    for s in staff:
        print(f"    • {s['full_name']} ({s['category']})")

    # Schritt 2: PubMed-Suche
    print(f"\n[2/4] PubMed-Suche für {len(staff)} Mitarbeiter...")
    all_articles = []
    for i, person in enumerate(staff, 1):
        print(f"  [{i}/{len(staff)}] {person['clean_name']}...", end=" ", flush=True)
        articles = search_pubmed(person, start_date_pm, end_date_pm)
        print(f"→ {len(articles)} Treffer")
        all_articles.extend(articles)

    print(f"  → {len(all_articles)} PubMed-Artikel insgesamt (vor Duplikat-Bereinigung)")

    # Schritt 3: Google Scholar (optional)
    if not args.no_scholar:
        print(f"\n[3/4] Google Scholar-Suche für {len(staff)} Mitarbeiter...")
        print("  (Dies kann dauern wegen Rate-Limiting...)")
        scholar_count = 0
        for i, person in enumerate(staff, 1):
            print(f"  [{i}/{len(staff)}] {person['clean_name']}...", end=" ", flush=True)
            articles = search_scholar(person, start_year, end_year)
            print(f"→ {len(articles)} Treffer")
            all_articles.extend(articles)
            scholar_count += len(articles)
        print(f"  → {scholar_count} Google Scholar-Artikel insgesamt")
    else:
        print("\n[3/4] Google Scholar übersprungen (--no-scholar)")

    # Schritt 4: Duplikat-Bereinigung und Word-Export
    print(f"\n[4/4] Duplikat-Bereinigung und Word-Export...")
    print(f"  {len(all_articles)} Artikel vor Bereinigung")
    unique_articles = deduplicate_articles(all_articles)
    print(f"  {len(unique_articles)} Artikel nach Bereinigung")

    create_word_document(
        unique_articles, staff, args.start_date, args.end_date, args.output
    )

    print(f"\n{'=' * 70}")
    print(f"FERTIG! Dokument gespeichert: {args.output}")
    print(f"  {len(unique_articles)} eindeutige Publikationen")
    print(f"  {len(staff)} Mitarbeiter durchsucht")
    print(f"{'=' * 70}")


if __name__ == "__main__":
    # Windows Konsolen-Encoding fix
    import io
    if sys.stdout.encoding != "utf-8":
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
        sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")
    main()
